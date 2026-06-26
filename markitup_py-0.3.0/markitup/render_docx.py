"""IR + Theme -> .docx via python-docx.

The renderer is mechanical: it never *decides* styling, it only reads tokens
from the Theme. Heading sizes come from the computed modular scale; colors,
fonts, margins, and the optional watermark all come from theme tokens.
"""
from __future__ import annotations

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from . import ir

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


# --- public API -------------------------------------------------------------
def render(doc_ir: ir.Document, theme, out_path: str) -> str:
    if theme.base_docx:
        # Template mode: the .docx owns the design. Open it, clear its body,
        # and map markdown onto its named styles. We do NOT override page or
        # heading styles — that is the template's job.
        document = Document(theme.base_docx)
        _clear_body(document)
        _ensure_aux_styles(document, theme)
    else:
        document = Document()
        _setup_page(document, theme)
        _build_styles(document, theme)
    if theme.watermark.enabled:
        _add_watermark(document.sections[0], theme.watermark, theme.fonts.heading)
    if theme.banner and theme.banner.text:
        _render_banner(document, theme)
    for block in doc_ir.children:
        _render_block(document, block, theme)
    document.save(out_path)
    return out_path


def _render_banner(document, theme):
    b = theme.banner
    p = document.add_paragraph()
    p.alignment = _ALIGN.get(b.align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(b.text)
    run.font.bold = b.bold
    run.font.size = Pt(b.size_pt or theme.type.base_size * 1.1)
    run.font.color.rgb = RGBColor.from_string(b.color)
    if b.bg:
        _set_para_shading(p, b.bg)
        # a little breathing room inside the colored bar
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)


def _clear_body(document):
    """Remove existing body content but keep the final sectPr (page setup,
    header/footer references)."""
    body = document.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


# --- page + styles ----------------------------------------------------------
def _setup_page(document, theme):
    sec = document.sections[0]
    if theme.page.size.upper() == "LETTER":
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    else:
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    m = Cm(theme.page.margin_cm)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m


def _build_styles(document, theme):
    normal = document.styles["Normal"]
    normal.font.name = theme.fonts.body
    normal.font.size = Pt(theme.type.base_size)
    normal.font.color.rgb = RGBColor.from_string(theme.colors.text)
    pf = normal.paragraph_format
    pf.line_spacing = theme.type.line_height
    pf.space_after = Pt(theme.type.base_size * 0.6)
    _force_font(normal, theme.fonts.body)

    for lvl in range(1, 7):
        st = document.styles[f"Heading {lvl}"]
        size = theme.type.heading_size(lvl)
        st.font.name = theme.fonts.heading
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.italic = False
        st.font.color.rgb = RGBColor.from_string(theme.colors.heading_color(lvl))
        hpf = st.paragraph_format
        hpf.space_before = Pt(size * 0.6)
        hpf.space_after = Pt(theme.type.base_size * 0.3)
        hpf.keep_with_next = True
        _force_font(st, theme.fonts.heading)

    _ensure_aux_styles(document, theme)


def _ensure_aux_styles(document, theme):
    """Code + Quote paragraph styles. Safe to call on a template that may or may
    not already define them."""
    names = [s.name for s in document.styles]
    if "MarkItUp Code" not in names:
        code = document.styles.add_style("MarkItUp Code", WD_STYLE_TYPE.PARAGRAPH)
        code.base_style = document.styles["Normal"]
        code.font.name = theme.fonts.mono
        code.font.size = Pt(theme.code_size)
        cpf = code.paragraph_format
        cpf.line_spacing = 1.0
        cpf.space_before = Pt(4)
        cpf.space_after = Pt(theme.type.base_size * 0.6)
        cpf.left_indent = Pt(6)
        _force_font(code, theme.fonts.mono)

    if "MarkItUp Quote" not in names:
        q = document.styles.add_style("MarkItUp Quote", WD_STYLE_TYPE.PARAGRAPH)
        q.base_style = document.styles["Normal"]
        q.font.italic = True
        q.paragraph_format.left_indent = Cm(0.8)


# --- block rendering --------------------------------------------------------
def _render_block(document, block, theme):
    if isinstance(block, ir.Heading):
        p = document.add_paragraph(style=f"Heading {min(block.level, 6)}")
        _add_inline(p, block.children, {}, theme)
    elif isinstance(block, ir.Paragraph):
        p = document.add_paragraph()
        _add_inline(p, block.children, {}, theme)
    elif isinstance(block, ir.CodeBlock):
        _render_code(document, block, theme)
    elif isinstance(block, ir.BlockQuote):
        for child in block.children:
            if isinstance(child, ir.Paragraph):
                p = document.add_paragraph(style="MarkItUp Quote")
                _left_border(p, theme.colors.accent)
                _add_inline(p, child.children, {"italic": True}, theme)
            else:
                _render_block(document, child, theme)
    elif isinstance(block, ir.ListNode):
        _render_list(document, block, theme, 1)
    elif isinstance(block, ir.Table):
        _render_table(document, block, theme)
    elif isinstance(block, ir.ThematicBreak):
        _hr(document.add_paragraph(), theme.colors.rule)


def _render_code(document, block, theme):
    p = document.add_paragraph(style="MarkItUp Code")
    _set_para_shading(p, theme.colors.code_bg)
    for i, line in enumerate(block.content.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = theme.fonts.mono
        r.font.size = Pt(theme.code_size)
        r.font.color.rgb = RGBColor.from_string(theme.colors.code_text)


def _render_list(document, lst, theme, level):
    for item in lst.items:
        first_done = False
        for b in item.children:
            if isinstance(b, ir.ListNode):
                _render_list(document, b, theme, min(level + 1, 3))
            elif isinstance(b, ir.Paragraph):
                p = document.add_paragraph(style=_list_style(lst.ordered, level))
                _add_inline(p, b.children, {}, theme)
                first_done = True
            else:
                _render_block(document, b, theme)
        if not first_done:
            document.add_paragraph(style=_list_style(lst.ordered, level))


def _list_style(ordered, level):
    base = "List Number" if ordered else "List Bullet"
    return base if level == 1 else f"{base} {level}"


def _render_table(document, t, theme):
    ncols = len(t.header.cells) if t.header else (max((len(r.cells) for r in t.rows), default=0))
    if ncols == 0:
        return
    tt = theme.table
    table = document.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    _set_table_width_full(table)
    _set_table_borders(table, tt.border_color, tt.border_width_pt)
    if t.header:
        cells = table.add_row().cells
        for j, cell in enumerate(t.header.cells[:ncols]):
            para = cells[j].paragraphs[0]
            para.alignment = _ALIGN.get(cell.align, WD_ALIGN_PARAGRAPH.LEFT)
            _add_inline(para, cell.children, {"bold": tt.header_bold, "color": tt.header_text}, theme)
            _set_cell_shading(cells[j], tt.header_bg)
    for row in t.rows:
        cells = table.add_row().cells
        for j, cell in enumerate(row.cells[:ncols]):
            para = cells[j].paragraphs[0]
            para.alignment = _ALIGN.get(cell.align, WD_ALIGN_PARAGRAPH.LEFT)
            _add_inline(para, cell.children, {"color": tt.body_text}, theme)


# --- inline rendering -------------------------------------------------------
def _add_inline(paragraph, nodes, fmt, theme):
    for n in nodes:
        if isinstance(n, ir.Text):
            _apply_run(paragraph.add_run(n.content), fmt, theme)
        elif isinstance(n, ir.LineBreak):
            paragraph.add_run().add_break()
        elif isinstance(n, ir.InlineCode):
            _apply_run(paragraph.add_run(n.content), {**fmt, "code": True}, theme)
        elif isinstance(n, ir.Strong):
            _add_inline(paragraph, n.children, {**fmt, "bold": True}, theme)
        elif isinstance(n, ir.Emphasis):
            _add_inline(paragraph, n.children, {**fmt, "italic": True}, theme)
        elif isinstance(n, ir.Strike):
            _add_inline(paragraph, n.children, {**fmt, "strike": True}, theme)
        elif isinstance(n, ir.Link):
            _add_hyperlink(paragraph, n.href, _plain(n.children), theme)
        elif isinstance(n, ir.Image):
            _apply_run(paragraph.add_run(n.alt or n.src), {**fmt, "italic": True}, theme)


def _apply_run(run, fmt, theme):
    if fmt.get("bold"):
        run.font.bold = True
    if fmt.get("italic"):
        run.font.italic = True
    if fmt.get("strike"):
        run.font.strike = True
    if fmt.get("color"):
        run.font.color.rgb = RGBColor.from_string(fmt["color"])
    if fmt.get("code"):
        run.font.name = theme.fonts.mono
        run.font.size = Pt(theme.code_size)
        run.font.color.rgb = RGBColor.from_string(theme.colors.code_text)
        run._element.get_or_add_rPr().append(_shd(theme.colors.code_bg))


def _add_hyperlink(paragraph, url, text, theme):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), theme.colors.link)
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# --- low-level OOXML helpers ------------------------------------------------
def _shd(fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    return shd


def _set_para_shading(p, fill):
    p._p.get_or_add_pPr().append(_shd(fill))


def _set_cell_shading(cell, fill):
    cell._tc.get_or_add_tcPr().append(_shd(fill))


def _set_table_borders(table, color, width_pt):
    """Uniform single borders. Word measures border size in eighths of a point,
    so 0.5pt -> sz=4."""
    sz = max(2, int(round(width_pt * 8)))
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _set_table_width_full(table):
    """Make the table span the text column and let Word distribute/wrap columns,
    so wide tables don't run past the margin."""
    tblPr = table._tbl.tblPr
    w = OxmlElement("w:tblW")
    w.set(qn("w:type"), "pct")
    w.set(qn("w:w"), "5000")  # 5000 fiftieths-of-percent == 100%
    tblPr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)


def _force_font(style, name):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)


def _hr(p, color):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _left_border(p, color):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    pbdr.append(left)
    pPr.append(pbdr)


_VML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office"'
)


def _add_watermark(section, wm, font):
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    if wm.image:
        p._p.append(parse_xml(_image_watermark_xml(header.part, wm)))
    else:
        p._p.append(parse_xml(_text_watermark_xml(wm, font)))


_VPOS = {"center": "center", "top": "top", "bottom": "bottom"}


def _text_watermark_xml(wm, font):
    text = (wm.text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    h = wm.width_pt * 0.25
    vpos = _VPOS.get(wm.position, "center")
    return (
        f'<w:r {_VML_NS}><w:pict>'
        f'<v:shape id="MarkItUpWatermark" o:spid="_x0000_s2049" type="#_x0000_t136" '
        f'style="position:absolute;margin-left:0;margin-top:0;width:{wm.width_pt}pt;height:{h}pt;'
        f'rotation:{wm.rotation};z-index:-1;mso-position-horizontal:center;'
        f'mso-position-horizontal-relative:margin;mso-position-vertical:{vpos};'
        'mso-position-vertical-relative:margin" '
        f'fillcolor="#{wm.color}" stroked="f">'
        f'<v:fill opacity="{wm.opacity}"/>'
        f'<v:textpath on="t" fitshape="t" string="{text}" '
        f'style="font-family:&quot;{font}&quot;;font-weight:bold"/>'
        '</v:shape></w:pict></w:r>'
    )


def _image_watermark_xml(header_part, wm):
    """Canonical Word image watermark: a VML shape in the header with washout
    (gain/blacklevel) so it sits faded behind the text."""
    rid, img = header_part.get_or_add_image(wm.image)
    try:
        dpi = img.horz_dpi or 96
        w_pt = img.px_width / dpi * 72.0
        h_pt = img.px_height / dpi * 72.0
        scale = wm.width_pt / w_pt if w_pt else 1.0
        w_pt, h_pt = w_pt * scale, h_pt * scale
    except Exception:
        w_pt, h_pt = wm.width_pt, wm.width_pt
    # washout gain/blacklevel emulate Word's "Washout" preset; opacity also set.
    return (
        f'<w:r {_VML_NS}><w:pict>'
        f'<v:shape id="MarkItUpWatermark" o:spid="_x0000_s2050" type="#_x0000_t75" '
        f'style="position:absolute;margin-left:0;margin-top:0;width:{w_pt:.1f}pt;height:{h_pt:.1f}pt;'
        f'rotation:{wm.rotation};z-index:-1;mso-position-horizontal:center;'
        f'mso-position-horizontal-relative:margin;mso-position-vertical:{_VPOS.get(wm.position, "center")};'
        'mso-position-vertical-relative:margin">'
        f'<v:imagedata r:id="{rid}" o:title="watermark" gain="19661f" blacklevel="22938f"/>'
        f'<v:fill opacity="{wm.opacity}"/>'
        '</v:shape></w:pict></w:r>'
    )


def _plain(nodes) -> str:
    out = []
    for n in nodes:
        if isinstance(n, (ir.Text, ir.InlineCode)):
            out.append(n.content)
        elif hasattr(n, "children"):
            out.append(_plain(n.children))
    return "".join(out)
